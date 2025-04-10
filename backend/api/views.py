from django.http import HttpResponse, JsonResponse
from django.views.decorators.csrf import csrf_exempt
from rest_framework.decorators import api_view
import docx
import json

@csrf_exempt  # ✅ ปิดการป้องกัน CSRF
@api_view(["POST"])
def generate_docx(request):
    try:
        data = json.loads(request.body.decode("utf-8"))
        project_title_th = data.get("project_title_th", "").strip()
        project_title_en = data.get("project_title_en", "").strip()
        author_th_1 = data.get("author_th_1", "").strip()
        author_th_2 = data.get("author_th_2", "").strip()
        academic_year = data.get("academic_year", "").strip()

        # if not all([project_title_th, project_title_en, author_th, author_th_2, academic_year]):
        #     return JsonResponse({"error": "กรุณากรอกข้อมูลให้ครบทุกช่อง"}, status=400)

        # ✅ สร้างเอกสาร Word
        doc = docx.Document()

        # ✅ ตั้งค่าฟอนต์ TH SarabunPSK
        from docx.shared import Pt
        from docx.oxml.ns import qn

        style = doc.styles["Normal"]
        style.font.name = "TH SarabunPSK"
        style.element.rPr.rFonts.set(qn("w:eastAsia"), "TH SarabunPSK")
        style.font.size = Pt(16)
        style.font.bold = True  # ✅ เพิ่มตัวหนาเป็นค่าเริ่มต้น
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
        # ✅ เพิ่มหน้าปก
        doc.add_paragraph("").alignment = 1  # จัดกึ่งกลาง
        title = doc.add_paragraph(project_title_th)  # ชื่อโครงการ (ไทย)
        title.alignment = 1  # จัดกึ่งกลาง
        doc.add_paragraph(project_title_en).alignment = 1  # ชื่อโครงการ (อังกฤษ)

        doc.add_paragraph("\n\n\n\n\n\n")
        doc.add_paragraph(f" {author_th_1}").alignment = 1  # ชื่อผู้จัดทำ (ไทย)
        doc.add_paragraph(f" {author_th_2}").alignment = 1  # ชื่อผู้จัดทำ (อังกฤษ)
        doc.add_paragraph("\n\n\n\n\n\n\n\n")
        doc.add_paragraph("ปริญญานิพนธ์นี้เป็นส่วนหนึ่งของการศึกษาตามหลักสูตรอุตสาหกรรมศาสตรบัณฑิต").alignment = 1
        doc.add_paragraph("สาขาวิชาเทคโนโลยีสารสนเทศ ภาควิชาเทคโนโลยีสารสนเทศ").alignment = 1
        doc.add_paragraph("คณะเทคโนโลยีและการจัดการอุตสาหกรรม").alignment = 1
        doc.add_paragraph("มหาวิทยาลัยเทคโนโลยีพระจอมเกล้าพระนครเหนือ").alignment = 1
        doc.add_paragraph(f"ปีการศึกษา {academic_year}").alignment = 1  # ปีการศึกษา
        doc.add_paragraph("ลิขสิทธิ์ของมหาวิทยาลัยเทคโนโลยีพระจอมเกล้าพระนครเหนือ").alignment = 1
        # doc.add_page_break()  # ✅ ขึ้นหน้าใหม่

        # ✅ ตั้งค่าการส่งไฟล์ Word กลับไปให้ React
        response = HttpResponse(content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        response["Content-Disposition"] = 'attachment; filename="project_cover.docx"'

        doc.save(response)
        return response

    except json.JSONDecodeError:
        return JsonResponse({"error": "Invalid JSON"}, status=400)
    



